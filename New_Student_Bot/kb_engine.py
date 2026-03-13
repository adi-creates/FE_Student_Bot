import json
import os
import pickle
import random
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Dict, List, Tuple

import numpy as np
from docx import Document


@dataclass
class FAQItem:
    question: str
    answer: str


class SimpleTfidfVectorizer:
    def __init__(self, ngram_range: Tuple[int, int] = (1, 2)):
        self.ngram_range = ngram_range
        self.vocabulary_: Dict[str, int] = {}
        self.idf_: np.ndarray = np.array([])

    def _tokenize(self, text: str) -> List[str]:
        return re.findall(r"[A-Za-z0-9]+", text.lower())

    def _ngrams(self, tokens: List[str]) -> List[str]:
        grams: List[str] = []
        min_n, max_n = self.ngram_range
        for n in range(min_n, max_n + 1):
            if len(tokens) < n:
                continue
            grams.extend(" ".join(tokens[i : i + n]) for i in range(len(tokens) - n + 1))
        return grams

    def fit(self, documents: List[str]) -> "SimpleTfidfVectorizer":
        doc_freq: Dict[str, int] = {}
        for doc in documents:
            tokens = self._tokenize(doc)
            grams = set(self._ngrams(tokens))
            for gram in grams:
                doc_freq[gram] = doc_freq.get(gram, 0) + 1

        self.vocabulary_ = {term: idx for idx, term in enumerate(sorted(doc_freq.keys()))}
        n_docs = max(len(documents), 1)
        self.idf_ = np.zeros(len(self.vocabulary_), dtype=np.float64)

        for term, idx in self.vocabulary_.items():
            df = doc_freq.get(term, 0)
            self.idf_[idx] = np.log((1 + n_docs) / (1 + df)) + 1.0

        return self

    def transform(self, documents: List[str]) -> np.ndarray:
        if not self.vocabulary_:
            raise ValueError("Vectorizer is not fitted.")

        matrix = np.zeros((len(documents), len(self.vocabulary_)), dtype=np.float64)

        for row_idx, doc in enumerate(documents):
            tokens = self._tokenize(doc)
            grams = self._ngrams(tokens)
            if not grams:
                continue

            term_counts: Dict[int, int] = {}
            for gram in grams:
                vocab_idx = self.vocabulary_.get(gram)
                if vocab_idx is not None:
                    term_counts[vocab_idx] = term_counts.get(vocab_idx, 0) + 1

            if not term_counts:
                continue

            max_tf = max(term_counts.values())
            for vocab_idx, count in term_counts.items():
                tf = count / max_tf
                matrix[row_idx, vocab_idx] = tf * self.idf_[vocab_idx]

            norm = np.linalg.norm(matrix[row_idx])
            if norm > 0:
                matrix[row_idx] = matrix[row_idx] / norm

        return matrix

    def fit_transform(self, documents: List[str]) -> np.ndarray:
        self.fit(documents)
        return self.transform(documents)


def _normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.strip())


def _line_is_question_candidate(line: str) -> bool:
    if not line:
        return False
    if len(line) > 260:
        return False
    if line.endswith("?"):
        return True
    # Matches formats like "1. What is ..." or "Q: ..."
    if re.match(r"^\d+[\).]\s+", line):
        return "?" in line or line.lower().startswith("what ")
    return False


def _extract_lines_from_docx(docx_path: str) -> List[str]:
    doc = Document(docx_path)
    lines: List[str] = []

    for paragraph in doc.paragraphs:
        text = _normalize_text(paragraph.text)
        if text:
            lines.append(text)

    # Also inspect table cells if FAQ content is tabular.
    for table in doc.tables:
        for row in table.rows:
            row_values = [_normalize_text(cell.text) for cell in row.cells]
            row_values = [value for value in row_values if value]
            if row_values:
                lines.append(" | ".join(row_values))

    return lines


def parse_faq_from_docx(docx_path: str) -> List[FAQItem]:
    lines = _extract_lines_from_docx(docx_path)
    faqs: List[FAQItem] = []

    current_question = ""
    answer_parts: List[str] = []
    in_answer = False

    def commit_current() -> None:
        nonlocal current_question, answer_parts, in_answer
        question = _normalize_text(current_question)
        answer = _normalize_text(" ".join(answer_parts))
        if question and answer:
            faqs.append(FAQItem(question=question, answer=answer))
        current_question = ""
        answer_parts = []
        in_answer = False

    for raw_line in lines:
        line = raw_line.strip()
        lower = line.lower()

        # Table style: "Question | Answer"
        if " | " in line:
            parts = [part.strip() for part in line.split("|") if part.strip()]
            if len(parts) >= 2:
                first = parts[0]
                second = " ".join(parts[1:])
                if first and second and ("?" in first or first.lower().startswith("q")):
                    commit_current()
                    faqs.append(FAQItem(question=first, answer=second))
                    continue

        q_match = re.match(r"^(q(?:uestion)?)\s*[:\-]\s*(.+)$", line, flags=re.IGNORECASE)
        a_match = re.match(r"^(a(?:nswer)?)\s*[:\-]\s*(.+)$", line, flags=re.IGNORECASE)

        if q_match:
            commit_current()
            current_question = q_match.group(2).strip()
            in_answer = False
            continue

        if a_match:
            if not current_question:
                continue
            in_answer = True
            answer_parts.append(a_match.group(2).strip())
            continue

        if _line_is_question_candidate(line):
            if current_question and answer_parts:
                commit_current()
            current_question = re.sub(r"^\d+[\).]\s+", "", line).strip()
            in_answer = False
            continue

        if current_question:
            in_answer = True
            answer_parts.append(line)

    commit_current()

    # Deduplicate by normalized question.
    unique: Dict[str, FAQItem] = {}
    for item in faqs:
        key = item.question.lower().strip()
        if key not in unique:
            unique[key] = item

    return list(unique.values())


def _validation_variants(question: str) -> List[str]:
    q = question.strip()
    no_punct = re.sub(r"[^A-Za-z0-9\s]", "", q)
    variants = [
        q,
        q.lower(),
        no_punct,
        f"Can you please tell me {no_punct}",
        f"I need help with: {no_punct}",
    ]
    return [variant for variant in variants if variant.strip()]


class KnowledgeBaseBot:
    def __init__(self, faqs: List[FAQItem], vectorizer: SimpleTfidfVectorizer, tfidf_matrix: np.ndarray):
        self.faqs = faqs
        self.vectorizer = vectorizer
        if hasattr(tfidf_matrix, "toarray"):
            tfidf_matrix = tfidf_matrix.toarray()
        self.tfidf_matrix = np.asarray(tfidf_matrix, dtype=np.float64)

    @classmethod
    def train_from_docx(cls, docx_path: str) -> "KnowledgeBaseBot":
        faqs = parse_faq_from_docx(docx_path)
        if not faqs:
            raise ValueError("No FAQ entries could be extracted from the document.")

        vectorizer = SimpleTfidfVectorizer(ngram_range=(1, 2))
        questions = [item.question for item in faqs]
        tfidf_matrix = vectorizer.fit_transform(questions)
        return cls(faqs=faqs, vectorizer=vectorizer, tfidf_matrix=tfidf_matrix)

    def ask(self, user_question: str) -> Tuple[FAQItem, float]:
        user_question = user_question.strip()
        if not user_question:
            raise ValueError("Question must not be empty.")

        query_vec = self.vectorizer.transform([user_question])
        if hasattr(query_vec, "toarray"):
            query_vec = query_vec.toarray()
        query_vec = np.asarray(query_vec, dtype=np.float64).reshape(-1)
        similarities = np.dot(self.tfidf_matrix, query_vec)
        idx = int(np.argmax(similarities))
        score = float(similarities[idx])
        return self.faqs[idx], score

    def validate(self) -> Dict[str, object]:
        checks = []
        for item in self.faqs:
            for variant in _validation_variants(item.question):
                predicted, score = self.ask(variant)
                checks.append(
                    {
                        "query": variant,
                        "expected": item.question,
                        "predicted": predicted.question,
                        "confidence": round(score, 4),
                        "correct": predicted.question == item.question,
                    }
                )

        total = len(checks)
        correct = sum(1 for entry in checks if entry["correct"])
        accuracy = correct / total if total else 0.0
        return {
            "faq_count": len(self.faqs),
            "validation_queries": total,
            "top1_accuracy": round(accuracy, 4),
            "sample": random.sample(checks, k=min(10, len(checks))) if checks else [],
        }

    def save(self, output_dir: str) -> None:
        path = Path(output_dir)
        path.mkdir(parents=True, exist_ok=True)

        with open(path / "faqs.json", "w", encoding="utf-8") as f:
            json.dump([asdict(item) for item in self.faqs], f, indent=2)

        with open(path / "model.pkl", "wb") as f:
            pickle.dump(
                {
                    "vectorizer": self.vectorizer,
                    "tfidf_matrix": self.tfidf_matrix,
                    "faqs": self.faqs,
                },
                f,
            )

        validation_report = self.validate()
        with open(path / "validation_report.json", "w", encoding="utf-8") as f:
            json.dump(validation_report, f, indent=2)

    @classmethod
    def load(cls, model_path: str) -> "KnowledgeBaseBot":
        if not os.path.exists(model_path):
            raise FileNotFoundError(f"Model artifact not found: {model_path}")

        with open(model_path, "rb") as f:
            payload = pickle.load(f)

        return cls(
            faqs=payload["faqs"],
            vectorizer=payload["vectorizer"],
            tfidf_matrix=payload["tfidf_matrix"],
        )
